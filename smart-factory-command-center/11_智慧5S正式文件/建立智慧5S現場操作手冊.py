#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
化新精密有限公司｜智慧 5S 現場操作手冊 Word 產生器

以受控 Markdown 建立正式 Word，採用緊湊參考指南版型、手冊型封面、
固定 DXA 表格、真實 Word 清單、頁碼與結構稽核。
"""

from __future__ import annotations

import importlib.util as 匯入工具
import json as 資料格式
import sys as 系統
import zipfile as 壓縮檔
from pathlib import Path as 路徑
from typing import List as 清單型別, Sequence as 序列型別
from xml.etree import ElementTree as 元素樹

from docx import Document as 文件
from docx.enum.text import WD_ALIGN_PARAGRAPH as 段落對齊, WD_BREAK as 換頁型態
from docx.oxml import OxmlElement as 開放元素
from docx.oxml.ns import qn as 合格名稱
from docx.shared import Inches as 英吋, Pt as 點


目前目錄 = 路徑(__file__).resolve().parent
來源檔案 = 目前目錄 / "智慧5S現場操作手冊_第1.0版.md"
輸出檔案 = 目前目錄 / "智慧5S現場操作手冊_第1.0版.docx"
稽核檔案 = 目前目錄 / "智慧5S現場操作手冊_第1.0版_結構稽核.json"
完整手冊產生器路徑 = 目前目錄 / "建立智慧5S完整操作手冊.py"


def 載入完整手冊工具():
    if not 完整手冊產生器路徑.exists():
        raise FileNotFoundError(f"找不到完整手冊產生器：{完整手冊產生器路徑}")
    規格資料 = 匯入工具.spec_from_file_location("智慧5S完整手冊工具", 完整手冊產生器路徑)
    if 規格資料 is None or 規格資料.loader is None:
        raise ImportError(f"無法載入完整手冊產生器：{完整手冊產生器路徑}")
    模組 = 匯入工具.module_from_spec(規格資料)
    規格資料.loader.exec_module(模組)
    return 模組


完整手冊工具 = 載入完整手冊工具()
規格 = 完整手冊工具.規格
管理工具 = 完整手冊工具.管理工具


def 選擇現場欄寬(表頭: 序列型別[str]) -> 清單型別[int]:
    欄數 = len(表頭)
    合併表頭 = "|".join(表頭)

    if 欄數 == 1:
        return [9360]

    if 欄數 == 2:
        if any(文字 in 合併表頭 for 文字 in ("合格照片", "可優先離線執行", "現場已完成")):
            return [4680, 4680]
        if "回報項目" in 合併表頭 or "必寫內容" in 合併表頭:
            return [2200, 7160]
        return [2500, 6860]

    if 欄數 == 3:
        if "畫面位置" in 合併表頭:
            return [1600, 3000, 4760]
        if "狀態" in 合併表頭:
            return [1500, 3100, 4760]
        if "寫法" in 合併表頭:
            return [1600, 4300, 3460]
        if "問題" in 合併表頭:
            return [2300, 3260, 3800]
        return [1900, 3600, 3860]

    if 欄數 == 4:
        if "管制項目" in 合併表頭:
            return [1500, 3180, 1500, 3180]
        if "編製" in 合併表頭 and "核准" in 合併表頭:
            return [2340, 2340, 2340, 2340]
        if "分數" in 合併表頭:
            return [900, 1800, 3000, 3660]
        if "我要做什麼" in 合併表頭:
            return [1600, 1900, 2500, 3360]
        if "原狀態" in 合併表頭:
            return [1400, 1600, 3100, 3260]
        return [1600, 2300, 2600, 2860]

    return 完整手冊工具.選擇手冊欄寬(表頭)


def 設定緊湊文件樣式(doc: 文件) -> None:
    規格.設定文件樣式(doc)
    樣式 = doc.styles

    正文 = 樣式["Normal"]
    正文.paragraph_format.space_before = 點(0)
    正文.paragraph_format.space_after = 點(6)
    正文.paragraph_format.line_spacing = 1.25

    標題一 = 樣式["Heading 1"]
    標題一.paragraph_format.space_before = 點(18)
    標題一.paragraph_format.space_after = 點(10)
    標題一.paragraph_format.line_spacing = 1.0

    標題二 = 樣式["Heading 2"]
    標題二.paragraph_format.space_before = 點(14)
    標題二.paragraph_format.space_after = 點(7)
    標題二.paragraph_format.line_spacing = 1.0

    標題三 = 樣式["Heading 3"]
    標題三.paragraph_format.space_before = 點(10)
    標題三.paragraph_format.space_after = 點(5)
    標題三.paragraph_format.line_spacing = 1.0


def 建立緊湊清單編號(doc: 文件, 類型: str) -> int:
    編號根 = doc.part.numbering_part.element
    抽象識別碼 = 規格.下一個識別碼(編號根, "w:abstractNum", "w:abstractNumId")
    清單識別碼 = 規格.下一個識別碼(編號根, "w:num", "w:numId")

    抽象編號 = 開放元素("w:abstractNum")
    抽象編號.set(合格名稱("w:abstractNumId"), str(抽象識別碼))
    多層型態 = 開放元素("w:multiLevelType")
    多層型態.set(合格名稱("w:val"), "singleLevel")
    抽象編號.append(多層型態)

    層級 = 開放元素("w:lvl")
    層級.set(合格名稱("w:ilvl"), "0")
    起始 = 開放元素("w:start")
    起始.set(合格名稱("w:val"), "1")
    編號格式 = 開放元素("w:numFmt")
    編號格式.set(合格名稱("w:val"), "bullet" if 類型 == "項目" else "decimal")
    層級文字 = 開放元素("w:lvlText")
    層級文字.set(合格名稱("w:val"), "●" if 類型 == "項目" else "%1.")
    層級對齊 = 開放元素("w:lvlJc")
    層級對齊.set(合格名稱("w:val"), "left")

    段落屬性 = 開放元素("w:pPr")
    定位集合 = 開放元素("w:tabs")
    定位點 = 開放元素("w:tab")
    定位點.set(合格名稱("w:val"), "num")
    定位點.set(合格名稱("w:pos"), "540")
    定位集合.append(定位點)
    縮排 = 開放元素("w:ind")
    縮排.set(合格名稱("w:left"), "540")
    縮排.set(合格名稱("w:hanging"), "270")
    間距 = 開放元素("w:spacing")
    間距.set(合格名稱("w:after"), "80")
    間距.set(合格名稱("w:line"), "300")
    間距.set(合格名稱("w:lineRule"), "auto")
    段落屬性.extend([定位集合, 縮排, 間距])
    層級.extend([起始, 編號格式, 層級文字, 層級對齊, 段落屬性])
    抽象編號.append(層級)
    編號根.append(抽象編號)

    清單編號 = 開放元素("w:num")
    清單編號.set(合格名稱("w:numId"), str(清單識別碼))
    抽象參照 = 開放元素("w:abstractNumId")
    抽象參照.set(合格名稱("w:val"), str(抽象識別碼))
    清單編號.append(抽象參照)
    編號根.append(清單編號)
    return 清單識別碼


def 新增緊湊清單項目(doc: 文件, 文字: str, 清單識別碼: int) -> None:
    段落 = doc.add_paragraph()
    段落.paragraph_format.space_before = 點(0)
    段落.paragraph_format.space_after = 點(4)
    段落.paragraph_format.line_spacing = 1.25
    規格.套用清單編號(段落, 清單識別碼)
    規格.加入格式化文字(段落, 文字)


def 新增緊湊核取項目(doc: 文件, 文字: str, 已完成: bool) -> None:
    段落 = doc.add_paragraph()
    段落.paragraph_format.left_indent = 英吋(0.375)
    段落.paragraph_format.space_before = 點(0)
    段落.paragraph_format.space_after = 點(4)
    段落.paragraph_format.line_spacing = 1.25
    標示 = "【完成】" if 已完成 else "【確認】"
    執行文字 = 段落.add_run(標示 + " ")
    規格.設定執行文字(
        執行文字,
        大小=10.5,
        顏色=規格.次深藍 if 已完成 else "7A5A00",
        粗體=True,
    )
    規格.加入格式化文字(段落, 文字)


原新增資料表 = 規格.新增資料表


def 新增不拆小表格(doc: 文件, 資料, *, 封面表格: bool = False) -> None:
    """讓七列以內的表格優先整體換頁，避免少量孤列落在下一頁。"""
    原新增資料表(doc, 資料, 封面表格=封面表格)
    if not 資料 or len(資料) > 7:
        return

    表格 = doc.tables[-1]
    for 資料列 in 表格.rows[:-1]:
        for 儲存格 in 資料列.cells:
            for 段落 in 儲存格.paragraphs:
                段落.paragraph_format.keep_with_next = True


# 套用緊湊參考指南的表格、清單與核取項目設計值。
規格.選擇欄寬 = 選擇現場欄寬
規格.淡灰 = 規格.淡藍灰
規格.新增資料表 = 新增不拆小表格
規格.建立清單編號 = 建立緊湊清單編號
規格.新增清單項目 = 新增緊湊清單項目
規格.新增核取項目 = 新增緊湊核取項目


def 設定現場頁首頁尾(區段) -> None:
    頁首 = 區段.header.paragraphs[0]
    規格.清除段落(頁首)
    頁首.alignment = 段落對齊.CENTER
    頁首.paragraph_format.space_after = 點(0)
    執行文字 = 頁首.add_run("化新精密有限公司｜智慧 5S 現場操作手冊｜第 1.0 版")
    規格.設定執行文字(執行文字, 大小=8.5, 顏色=規格.灰字)

    首頁頁首 = 區段.first_page_header.paragraphs[0]
    規格.清除段落(首頁頁首)

    for 頁尾物件 in (區段.footer, 區段.first_page_footer):
        頁尾 = 頁尾物件.paragraphs[0]
        規格.清除段落(頁尾)
        頁尾.alignment = 段落對齊.CENTER
        頁尾.paragraph_format.space_before = 點(0)
        頁尾.paragraph_format.space_after = 點(0)
        執行文字 = 頁尾.add_run("化新精密有限公司｜智慧 5S 現場操作手冊｜第 ")
        規格.設定執行文字(執行文字, 大小=8.5, 顏色=規格.灰字)
        規格.加入頁碼欄位(頁尾)
        執行文字 = 頁尾.add_run(" 頁")
        規格.設定執行文字(執行文字, 大小=8.5, 顏色=規格.灰字)


def 新增現場封面(doc: 文件, 前言: 序列型別[str]) -> None:
    文件資料, 核心要求 = 管理工具.解析封面資料(前言)

    公司 = doc.add_paragraph()
    公司.paragraph_format.space_before = 點(42)
    公司.paragraph_format.space_after = 點(18)
    公司.alignment = 段落對齊.CENTER
    執行文字 = 公司.add_run("化新精密有限公司｜正式現場文件")
    規格.設定執行文字(執行文字, 大小=11, 顏色=規格.金色, 粗體=True)

    類別 = doc.add_paragraph()
    類別.paragraph_format.space_before = 點(0)
    類別.paragraph_format.space_after = 點(8)
    類別.alignment = 段落對齊.CENTER
    執行文字 = 類別.add_run("現場人員隨手操作指南")
    規格.設定執行文字(執行文字, 大小=11, 顏色=規格.次深藍, 粗體=True)

    標題 = doc.add_paragraph()
    標題.paragraph_format.space_before = 點(0)
    標題.paragraph_format.space_after = 點(8)
    標題.alignment = 段落對齊.CENTER
    執行文字 = 標題.add_run("智慧 5S\n現場操作手冊")
    規格.設定執行文字(執行文字, 大小=30, 顏色=規格.深藍, 粗體=True)

    副標題 = doc.add_paragraph()
    副標題.paragraph_format.space_before = 點(0)
    副標題.paragraph_format.space_after = 點(24)
    副標題.alignment = 段落對齊.CENTER
    執行文字 = 副標題.add_run("開班、巡檢、改善、紅牌、同步與交班，一步一步照著做")
    規格.設定執行文字(執行文字, 大小=13, 顏色=規格.灰字)

    封面資料 = [
        ["管制項目", "內容", "管制項目", "內容"],
        ["文件編號", 文件資料.get("文件編號", "化新－5S－現場－001"), "文件版本", 文件資料.get("文件版本", "第 1.0 版")],
        ["對應系統", 文件資料.get("對應系統", "智慧 5S 管理平台第 1.0.2 版"), "文件狀態", 文件資料.get("文件狀態", "正式版（核准後生效）")],
        ["制定單位", 文件資料.get("制定單位", "製造部"), "制定日期", 文件資料.get("制定日期", "2026 年 8 月 9 日")],
    ]
    規格.新增資料表(doc, 封面資料, 封面表格=True)

    if 核心要求:
        規格.新增提示框(doc, 核心要求, 封面=True)

    核准標示 = doc.add_paragraph()
    核准標示.paragraph_format.space_before = 點(6)
    核准標示.paragraph_format.space_after = 點(5)
    核准標示.alignment = 段落對齊.CENTER
    執行文字 = 核准標示.add_run("核准與生效")
    規格.設定執行文字(執行文字, 大小=10.5, 顏色=規格.次深藍, 粗體=True)
    規格.新增資料表(doc, [
        ["編製", "審查", "核准", "生效日期"],
        ["製造部：____________", "單位主管：____________", "最高主管：____________", "______ 年 ____ 月 ____ 日"],
    ])

    提醒 = doc.add_paragraph()
    提醒.paragraph_format.space_before = 點(8)
    提醒.paragraph_format.space_after = 點(0)
    提醒.alignment = 段落對齊.CENTER
    執行文字 = 提醒.add_run("受控文件｜只用本人身分｜每次送出都確認單號、狀態與待同步筆數")
    規格.設定執行文字(執行文字, 大小=9.5, 顏色=規格.灰字, 粗體=True)
    提醒.add_run().add_break(換頁型態.PAGE)


def 建立文件() -> 文件:
    if not 來源檔案.exists():
        raise FileNotFoundError(f"找不到來源檔案：{來源檔案}")
    全部行 = 來源檔案.read_text(encoding="utf-8").splitlines()
    try:
        分隔位置 = 全部行.index("---")
    except ValueError as 錯誤:
        raise ValueError("來源文件缺少封面與本文分隔線 ---") from 錯誤

    doc = 文件()
    doc.core_properties.title = "化新精密有限公司｜智慧 5S 現場操作手冊"
    doc.core_properties.subject = "智慧 5S 管理平台第 1.0.2 版現場人員操作指南"
    doc.core_properties.author = "化新精密有限公司"
    doc.core_properties.last_modified_by = "化新精密有限公司"
    doc.core_properties.keywords = "智慧5S,現場操作,巡檢,改善,紅牌,同步,交班"
    doc.core_properties.comments = "由受控 Markdown 來源自動產生；畫面、程序與系統版本同步。"

    設定緊湊文件樣式(doc)
    區段 = doc.sections[0]
    規格.設定頁面(區段)
    設定現場頁首頁尾(區段)
    新增現場封面(doc, 全部行[:分隔位置])
    規格.轉換本文(doc, 全部行[分隔位置 + 1:])
    return doc


def 擷取全文(檔案: 路徑) -> str:
    命名空間 = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with 壓縮檔.ZipFile(檔案) as 封裝:
        文件根 = 元素樹.fromstring(封裝.read("word/document.xml"))
    return "".join(節點.text or "" for 節點 in 文件根.findall(".//w:t", 命名空間))


def 尋找樣式(樣式根, 樣式識別碼: str):
    for 樣式 in 樣式根.findall(合格名稱("w:style")):
        if 樣式.get(合格名稱("w:styleId")) == 樣式識別碼:
            return 樣式
    return None


def 稽核緊湊版型(檔案: 路徑) -> 清單型別[str]:
    命名空間 = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    錯誤: 清單型別[str] = []
    with 壓縮檔.ZipFile(檔案) as 封裝:
        文件根 = 元素樹.fromstring(封裝.read("word/document.xml"))
        樣式根 = 元素樹.fromstring(封裝.read("word/styles.xml"))
        編號根 = 元素樹.fromstring(封裝.read("word/numbering.xml"))

    樣式期望 = {
        "Normal": {"大小": "22", "顏色": "000000", "前": "0", "後": "120", "行距": "300"},
        "Heading1": {"大小": "32", "顏色": "2E74B5", "前": "360", "後": "200"},
        "Heading2": {"大小": "26", "顏色": "2E74B5", "前": "280", "後": "140"},
        "Heading3": {"大小": "24", "顏色": "1F4D78", "前": "200", "後": "100"},
    }
    for 樣式識別碼, 期望 in 樣式期望.items():
        樣式 = 尋找樣式(樣式根, 樣式識別碼)
        if 樣式 is None:
            錯誤.append(f"缺少緊湊版型樣式：{樣式識別碼}")
            continue
        大小 = 樣式.find("./w:rPr/w:sz", 命名空間)
        顏色 = 樣式.find("./w:rPr/w:color", 命名空間)
        間距 = 樣式.find("./w:pPr/w:spacing", 命名空間)
        if 大小 is None or 大小.get(合格名稱("w:val")) != 期望["大小"]:
            錯誤.append(f"{樣式識別碼} 字級不符合緊湊版型")
        if 顏色 is None or 顏色.get(合格名稱("w:val"), "").upper() != 期望["顏色"]:
            錯誤.append(f"{樣式識別碼} 顏色不符合緊湊版型")
        if 間距 is None:
            錯誤.append(f"{樣式識別碼} 缺少明確段落間距")
            continue
        if 間距.get(合格名稱("w:before"), "0") != 期望["前"]:
            錯誤.append(f"{樣式識別碼} 段前不符合緊湊版型")
        if 間距.get(合格名稱("w:after"), "0") != 期望["後"]:
            錯誤.append(f"{樣式識別碼} 段後不符合緊湊版型")
        if "行距" in 期望 and 間距.get(合格名稱("w:line")) != 期望["行距"]:
            錯誤.append(f"{樣式識別碼} 行距不符合緊湊版型")

    清單參照 = {
        節點.get(合格名稱("w:val"))
        for 節點 in 文件根.findall(".//w:body/w:p/w:pPr/w:numPr/w:numId", 命名空間)
        if 節點.get(合格名稱("w:val"))
    }
    編號對抽象 = {}
    for 編號 in 編號根.findall("./w:num", 命名空間):
        參照 = 編號.find("./w:abstractNumId", 命名空間)
        if 參照 is not None:
            編號對抽象[編號.get(合格名稱("w:numId"))] = 參照.get(合格名稱("w:val"))
    抽象對層級 = {
        抽象.get(合格名稱("w:abstractNumId")): 抽象.find("./w:lvl", 命名空間)
        for 抽象 in 編號根.findall("./w:abstractNum", 命名空間)
    }
    for 清單識別碼 in 清單參照:
        層級 = 抽象對層級.get(編號對抽象.get(清單識別碼))
        縮排 = 層級.find("./w:pPr/w:ind", 命名空間) if 層級 is not None else None
        間距 = 層級.find("./w:pPr/w:spacing", 命名空間) if 層級 is not None else None
        定位點 = 層級.find("./w:pPr/w:tabs/w:tab", 命名空間) if 層級 is not None else None
        if (
            縮排 is None
            or 縮排.get(合格名稱("w:left")) != "540"
            or 縮排.get(合格名稱("w:hanging")) != "270"
            or 定位點 is None
            or 定位點.get(合格名稱("w:pos")) != "540"
            or 間距 is None
            or 間距.get(合格名稱("w:after")) != "80"
            or 間距.get(合格名稱("w:line")) != "300"
        ):
            錯誤.append(f"清單 {清單識別碼} 不符合緊湊版型縮排或間距")

    for 表格序號, 表格 in enumerate(文件根.findall(".//w:tbl", 命名空間), start=1):
        列 = 表格.findall("./w:tr", 命名空間)
        if len(列) <= 1:
            continue
        第一格 = 列[0].find("./w:tc", 命名空間)
        底色 = 第一格.find("./w:tcPr/w:shd", 命名空間) if 第一格 is not None else None
        if 底色 is None or 底色.get(合格名稱("w:fill"), "").upper() != "E8EEF5":
            錯誤.append(f"第 {表格序號} 個資料表表頭底色不符合緊湊版型")

    return 錯誤


def 稽核現場手冊(檔案: 路徑) -> dict:
    報告 = 規格.稽核結構(檔案)
    報告["版型"] = "緊湊參考指南"
    報告["封面"] = "手冊型封面"
    報告["樣式基準"] = {
        "正文": "11 點、段後 6 點、1.25 倍行距",
        "第一層標題": "16 點、段前 18 點、段後 10 點",
        "第二層標題": "13 點、段前 14 點、段後 7 點",
        "第三層標題": "12 點、段前 10 點、段後 5 點",
        "清單": "標記 0.187 英吋、文字 0.375 英吋、段後 4 點、1.25 倍行距",
        "表格": "9360 DXA、縮排 120 DXA、表頭淡藍灰",
    }
    報告["結構錯誤"].extend(稽核緊湊版型(檔案))

    全文 = 擷取全文(檔案)
    必備文字 = (
        "化新－5S－現場－001",
        "開班六十秒檢查",
        "登入智慧 5S",
        "立即開始巡檢",
        "送出巡檢",
        "低於 5 分",
        "開始改善",
        "完成並送驗",
        "＋ 新增物品盤點",
        "開始處理",
        "立即同步",
        "清除快取並更新",
        "待同步 0",
        "異常回報六要素",
        "收班七步驟",
        "文件結束｜",
    )
    for 文字 in 必備文字:
        if 文字 not in 全文:
            報告["結構錯誤"].append(f"缺少現場手冊基準：{文字}")
    if 報告.get("表格數", 0) < 25:
        報告["結構錯誤"].append("現場手冊表格數不足 25，疑似內容未完整輸出")
    if 報告.get("真實編號定義數", 0) < 20:
        報告["結構錯誤"].append("現場手冊真實編號定義不足 20，疑似步驟未完整輸出")

    報告["結構錯誤數"] = len(報告["結構錯誤"])
    報告["結果"] = "通過" if not 報告["結構錯誤"] else "失敗"
    return 報告


def 主程式() -> int:
    doc = 建立文件()
    doc.save(輸出檔案)
    報告 = 稽核現場手冊(輸出檔案)
    稽核檔案.write_text(資料格式.dumps(報告, ensure_ascii=False, indent=2), encoding="utf-8")
    print(資料格式.dumps(報告, ensure_ascii=False, indent=2))
    print(f"已產生：{輸出檔案}")
    return 1 if 報告["結構錯誤數"] else 0


if __name__ == "__main__":
    系統.exit(主程式())
