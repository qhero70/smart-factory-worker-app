#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
化新精密有限公司｜智慧 5S 系統管理手冊 Word 產生器

以受控 Markdown 建立正式 Word，沿用智慧 5S 正式手冊的緊湊參考指南版型，
並針對唯一架構、帳號、資料庫、PWA、GAS、LINE、離線、備份與復原進行結構稽核。
"""

from __future__ import annotations

import importlib.util as 匯入工具
import json as 資料格式
import sys as 系統
from pathlib import Path as 路徑
from typing import List as 清單型別, Sequence as 序列型別

from docx import Document as 文件
from docx.enum.text import WD_ALIGN_PARAGRAPH as 段落對齊, WD_BREAK as 換頁型態
from docx.shared import Pt as 點


目前目錄 = 路徑(__file__).resolve().parent
來源檔案 = 目前目錄 / "智慧5S系統管理手冊_第1.0版.md"
輸出檔案 = 目前目錄 / "智慧5S系統管理手冊_第1.0版.docx"
稽核檔案 = 目前目錄 / "智慧5S系統管理手冊_第1.0版_結構稽核.json"
現場手冊產生器路徑 = 目前目錄 / "建立智慧5S現場操作手冊.py"


def 載入現場手冊工具():
    if not 現場手冊產生器路徑.exists():
        raise FileNotFoundError(f"找不到現場手冊產生器：{現場手冊產生器路徑}")
    規格資料 = 匯入工具.spec_from_file_location("智慧5S現場手冊工具", 現場手冊產生器路徑)
    if 規格資料 is None or 規格資料.loader is None:
        raise ImportError(f"無法載入現場手冊產生器：{現場手冊產生器路徑}")
    模組 = 匯入工具.module_from_spec(規格資料)
    規格資料.loader.exec_module(模組)
    return 模組


現場手冊工具 = 載入現場手冊工具()
規格 = 現場手冊工具.規格
管理工具 = 現場手冊工具.管理工具


def 選擇系統欄寬(表頭: 序列型別[str]) -> 清單型別[int]:
    """使用固定 9360 DXA，讓識別碼、欄位規格與故障處理欄保有閱讀寬度。"""
    欄數 = len(表頭)
    合併表頭 = "|".join(表頭)

    if 欄數 == 1:
        return [9360]
    if 欄數 == 2:
        if any(文字 in 合併表頭 for 文字 in ("正式欄位順序", "填寫內容", "最低內容")):
            return [2200, 7160]
        if any(文字 in 合併表頭 for 文字 in ("回復依據", "必須一起查核", "初始化值")):
            return [2800, 6560]
        return [2500, 6860]
    if 欄數 == 3:
        if "第一查核點" in 合併表頭:
            return [1700, 3000, 4660]
        if "正常" in 合併表頭 or "通過訊號" in 合併表頭:
            return [2100, 2300, 4960]
        if "正式值" in 合併表頭 or "所在專案" in 合併表頭:
            return [1900, 3300, 4160]
        return [1900, 3200, 4260]
    if 欄數 == 4:
        if "管制項目" in 合併表頭:
            return [1500, 3180, 1500, 3180]
        if "編製" in 合併表頭 and "核准" in 合併表頭:
            return [2340, 2340, 2340, 2340]
        if "角色" in 合併表頭:
            return [1450, 2600, 2600, 2710]
        if "分頁" in 合併表頭:
            return [2000, 2800, 1800, 2760]
        return [1500, 2500, 2500, 2860]
    return 現場手冊工具.選擇現場欄寬(表頭)


# 沿用 compact_reference_guide 的樣式、真實清單與固定表格，只調整系統文件欄寬。
規格.選擇欄寬 = 選擇系統欄寬


def 設定系統頁首頁尾(區段) -> None:
    頁首 = 區段.header.paragraphs[0]
    規格.清除段落(頁首)
    頁首.alignment = 段落對齊.CENTER
    頁首.paragraph_format.space_after = 點(0)
    執行文字 = 頁首.add_run("化新精密有限公司｜智慧 5S 系統管理手冊｜第 1.0 版")
    規格.設定執行文字(執行文字, 大小=8.5, 顏色=規格.灰字)

    首頁頁首 = 區段.first_page_header.paragraphs[0]
    規格.清除段落(首頁頁首)

    for 頁尾物件 in (區段.footer, 區段.first_page_footer):
        頁尾 = 頁尾物件.paragraphs[0]
        規格.清除段落(頁尾)
        頁尾.alignment = 段落對齊.CENTER
        頁尾.paragraph_format.space_before = 點(0)
        頁尾.paragraph_format.space_after = 點(0)
        執行文字 = 頁尾.add_run("化新精密有限公司｜智慧 5S 系統管理手冊｜第 ")
        規格.設定執行文字(執行文字, 大小=8.5, 顏色=規格.灰字)
        規格.加入頁碼欄位(頁尾)
        執行文字 = 頁尾.add_run(" 頁")
        規格.設定執行文字(執行文字, 大小=8.5, 顏色=規格.灰字)


def 新增系統封面(doc: 文件, 前言: 序列型別[str]) -> None:
    文件資料, 核心要求 = 管理工具.解析封面資料(前言)

    公司 = doc.add_paragraph()
    公司.paragraph_format.space_before = 點(42)
    公司.paragraph_format.space_after = 點(18)
    公司.alignment = 段落對齊.CENTER
    執行文字 = 公司.add_run("化新精密有限公司｜正式系統文件")
    規格.設定執行文字(執行文字, 大小=11, 顏色=規格.金色, 粗體=True)

    類別 = doc.add_paragraph()
    類別.paragraph_format.space_after = 點(8)
    類別.alignment = 段落對齊.CENTER
    執行文字 = 類別.add_run("系統管理、維運、發布與復原指南")
    規格.設定執行文字(執行文字, 大小=11, 顏色=規格.次深藍, 粗體=True)

    標題 = doc.add_paragraph()
    標題.paragraph_format.space_after = 點(8)
    標題.alignment = 段落對齊.CENTER
    執行文字 = 標題.add_run("智慧 5S\n系統管理手冊")
    規格.設定執行文字(執行文字, 大小=30, 顏色=規格.深藍, 粗體=True)

    副標題 = doc.add_paragraph()
    副標題.paragraph_format.space_after = 點(24)
    副標題.alignment = 段落對齊.CENTER
    執行文字 = 副標題.add_run("唯一架構、帳號、資料庫、PWA、GAS、LINE、離線、備份與復原")
    規格.設定執行文字(執行文字, 大小=13, 顏色=規格.灰字)

    封面資料 = [
        ["管制項目", "內容", "管制項目", "內容"],
        ["文件編號", 文件資料.get("文件編號", "化新－5S－系統－001"), "文件版本", 文件資料.get("文件版本", "第 1.0 版")],
        ["對應系統", 文件資料.get("對應系統", "智慧 5S 管理平台第 1.0.2 版"), "文件狀態", 文件資料.get("文件狀態", "正式版（核准後生效）")],
        ["制定單位", 文件資料.get("制定單位", "製造部／資訊管理"), "制定日期", 文件資料.get("制定日期", "2026 年 8 月 10 日")],
    ]
    規格.新增資料表(doc, 封面資料, 封面表格=True)
    if 核心要求:
        規格.新增提示框(doc, 核心要求, 封面=True)

    核准標示 = doc.add_paragraph()
    核准標示.paragraph_format.space_before = 點(6)
    核准標示.paragraph_format.space_after = 點(5)
    核准標示.alignment = 段落對齊.CENTER
    執行文字 = 核准標示.add_run("核准與生效")
    規格.設定執行文字(執行文字, 大小=10.5, 顏色=規格.次深藍, 粗體=True)
    規格.新增資料表(doc, [
        ["編製", "審查", "核准", "生效日期"],
        ["製造部／資訊：____________", "系統管理：____________", "最高主管：____________", "______ 年 ____ 月 ____ 日"],
    ])

    提醒 = doc.add_paragraph()
    提醒.paragraph_format.space_before = 點(8)
    提醒.paragraph_format.space_after = 點(0)
    提醒.alignment = 段落對齊.CENTER
    執行文字 = 提醒.add_run("維運口訣：先確認唯一來源、先備份、最小變更、完整驗收、可回復、留紀錄。")
    規格.設定執行文字(執行文字, 大小=9.5, 顏色=規格.灰字, 粗體=True)
    提醒.add_run().add_break(換頁型態.PAGE)


def 建立文件() -> 文件:
    if not 來源檔案.exists():
        raise FileNotFoundError(f"找不到來源檔案：{來源檔案}")
    全部行 = 來源檔案.read_text(encoding="utf-8").splitlines()
    try:
        分隔位置 = 全部行.index("---")
    except ValueError as 錯誤:
        raise ValueError("來源文件缺少封面與本文分隔線 ---") from 錯誤

    doc = 文件()
    doc.core_properties.title = "化新精密有限公司｜智慧 5S 系統管理手冊"
    doc.core_properties.subject = "智慧 5S 管理平台第 1.0.2 版系統管理、維運、發布與復原指南"
    doc.core_properties.author = "化新精密有限公司"
    doc.core_properties.last_modified_by = "化新精密有限公司"
    doc.core_properties.keywords = "智慧5S,系統管理,PWA,GAS,試算表,LINE,離線同步,備份,復原,發布"
    doc.core_properties.comments = "由受控 Markdown 來源自動產生；架構、程序與系統版本同步。"

    現場手冊工具.設定緊湊文件樣式(doc)
    區段 = doc.sections[0]
    規格.設定頁面(區段)
    設定系統頁首頁尾(區段)
    新增系統封面(doc, 全部行[:分隔位置])
    規格.轉換本文(doc, 全部行[分隔位置 + 1:])
    return doc


def 稽核系統手冊(檔案: 路徑) -> dict:
    報告 = 規格.稽核結構(檔案)
    報告["版型"] = "緊湊參考指南"
    報告["封面"] = "編輯型系統管理手冊封面"
    報告["樣式基準"] = {
        "正文": "11 點、段後 6 點、1.25 倍行距",
        "第一層標題": "16 點、段前 18 點、段後 10 點",
        "第二層標題": "13 點、段前 14 點、段後 7 點",
        "第三層標題": "12 點、段前 10 點、段後 5 點",
        "清單": "標記 0.187 英吋、文字 0.375 英吋、段後 4 點、1.25 倍行距",
        "表格": "9360 DXA、縮排 120 DXA、表頭淡藍灰",
    }
    報告["結構錯誤"].extend(現場手冊工具.稽核緊湊版型(檔案))

    全文 = 現場手冊工具.擷取全文(檔案)
    必備文字 = (
        "化新－5S－系統－001",
        "正式架構與唯一來源",
        "帳號、角色與存取安全",
        "十五張必要分頁",
        "安全初始化程序",
        "交易資料與資料修正",
        "PWA 設定、版本與發布",
        "Apps Script 後端管理",
        "唯一 LINE Bot 與通知管理",
        "離線同步、快取與裝置管理",
        "照片、容量與個資管理",
        "備份、復原與營運持續",
        "故障診斷與事件處理",
        "資安與資料治理強化",
        "系統管理員快速操作卡",
        "文件結束｜",
    )
    for 文字 in 必備文字:
        if 文字 not in 全文:
            報告["結構錯誤"].append(f"缺少系統管理手冊基準：{文字}")
    if 報告.get("表格數", 0) < 45:
        報告["結構錯誤"].append("系統管理手冊表格數不足 45，疑似內容未完整輸出")
    if 報告.get("真實編號定義數", 0) < 35:
        報告["結構錯誤"].append("系統管理手冊真實編號定義不足 35，疑似程序未完整輸出")

    報告["結構錯誤數"] = len(報告["結構錯誤"])
    報告["結果"] = "通過" if not 報告["結構錯誤"] else "失敗"
    return 報告


def 主程式() -> int:
    doc = 建立文件()
    doc.save(輸出檔案)
    報告 = 稽核系統手冊(輸出檔案)
    稽核檔案.write_text(資料格式.dumps(報告, ensure_ascii=False, indent=2), encoding="utf-8")
    print(資料格式.dumps(報告, ensure_ascii=False, indent=2))
    print(f"已產生：{輸出檔案}")
    return 1 if 報告["結構錯誤數"] else 0


if __name__ == "__main__":
    系統.exit(主程式())
