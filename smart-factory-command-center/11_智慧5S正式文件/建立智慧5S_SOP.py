#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
化新精密有限公司｜智慧 5S 標準作業程序 Word 產生器

以受控 Markdown 建立正式 Word，沿用智慧 5S 文件設計系統、固定 DXA 表格、
真實 Word 清單、頁碼與結構稽核。文件文字只維護在智慧5S_SOP_v1.0.md。
"""

from __future__ import annotations

import importlib.util
import json
import sys
import zipfile
from pathlib import Path
from typing import List, Sequence
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


目前目錄 = Path(__file__).resolve().parent
來源檔案 = 目前目錄 / "智慧5S_SOP_v1.0.md"
輸出檔案 = 目前目錄 / "智慧5S_SOP_v1.0.docx"
稽核檔案 = 目前目錄 / "智慧5S_SOP_v1.0_結構稽核.json"
管理辦法產生器路徑 = 目前目錄 / "建立智慧5S管理辦法.py"


def 載入管理辦法工具():
    if not 管理辦法產生器路徑.exists():
        raise FileNotFoundError(f"找不到管理辦法產生器：{管理辦法產生器路徑}")
    spec = importlib.util.spec_from_file_location("智慧5S管理辦法工具", 管理辦法產生器路徑)
    if spec is None or spec.loader is None:
        raise ImportError(f"無法載入管理辦法產生器：{管理辦法產生器路徑}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


管理工具 = 載入管理辦法工具()
規格 = 管理工具.規格
基礎欄寬 = 管理工具.原選擇欄寬


def 選擇SOP欄寬(headers: Sequence[str]) -> List[int]:
    數量 = len(headers)
    joined = "|".join(headers)

    if 數量 == 2:
        return [2200, 7160]

    if 數量 == 3:
        if "合格做法" in joined:
            return [1500, 3930, 3930]
        if "完成訊號" in joined:
            return [1500, 3500, 4360]
        if "自動建立" in joined:
            return [1800, 4200, 3360]
        if "必留項目" in joined:
            return [1800, 7560]
        if "至少應保留" in joined:
            return [1800, 4200, 3360]
        return [1800, 3780, 3780]

    if 數量 == 4:
        if "管制項目" in joined:
            return [1500, 3180, 1500, 3180]
        if "編製" in joined and "核准" in joined:
            return [2340, 2340, 2340, 2340]
        if "版次" in joined:
            return [1000, 1500, 5060, 1800]
        if "程序" in joined and "正式輸出" in joined:
            return [1200, 2400, 2400, 3360]
        if "對象" in joined and "必備內容" in joined:
            return [1300, 4760, 1700, 1600]
        if "畫面按鈕" in joined:
            return [1500, 1700, 3300, 2860]
        if "判定" in joined and "必填／必做" in joined:
            return [1200, 2900, 2900, 2360]
        if "場景" in joined and "去重鍵" in joined:
            return [1800, 3400, 2200, 1960]
        if "主題" in joined and "管理問題" in joined:
            return [1500, 2800, 2500, 2560]
        if "順序" in joined and "允許動作" in joined:
            return [800, 1700, 3800, 3060]
        return 基礎欄寬(headers)

    if 數量 == 5:
        if "原狀態" in joined and "最低證據" in joined:
            return [1000, 1500, 1850, 1500, 3510]
        if "現場人員" in joined and "系統管理員" in joined:
            return [2400, 1700, 1800, 1500, 1960]
        return [1400, 1800, 2200, 1700, 2260]

    return 基礎欄寬(headers)


規格.選擇欄寬 = 選擇SOP欄寬


def 設定SOP頁首頁尾(section) -> None:
    頁首 = section.header.paragraphs[0]
    規格.清除段落(頁首)
    頁首.alignment = WD_ALIGN_PARAGRAPH.CENTER
    頁首.paragraph_format.space_after = Pt(0)
    run = 頁首.add_run("化新精密有限公司｜智慧 5S 管理平台 v1.0.2")
    規格.設定執行文字(run, 大小=8.5, 顏色=規格.灰字)

    首頁頁首 = section.first_page_header.paragraphs[0]
    規格.清除段落(首頁頁首)

    for footer in (section.footer, section.first_page_footer):
        頁尾 = footer.paragraphs[0]
        規格.清除段落(頁尾)
        頁尾.alignment = WD_ALIGN_PARAGRAPH.CENTER
        頁尾.paragraph_format.space_before = Pt(0)
        頁尾.paragraph_format.space_after = Pt(0)
        run = 頁尾.add_run("化新精密有限公司｜智慧 5S 標準作業程序｜第 ")
        規格.設定執行文字(run, 大小=8.5, 顏色=規格.灰字)
        規格.加入頁碼欄位(頁尾)
        run = 頁尾.add_run(" 頁")
        規格.設定執行文字(run, 大小=8.5, 顏色=規格.灰字)


def 新增SOP封面(doc: Document, 前言: Sequence[str]) -> None:
    metadata, 核心要求 = 管理工具.解析封面資料(前言)

    masthead = doc.add_paragraph()
    masthead.paragraph_format.space_before = Pt(18)
    masthead.paragraph_format.space_after = Pt(14)
    run = masthead.add_run("化新精密有限公司｜受控標準作業程序")
    規格.設定執行文字(run, 大小=11, 顏色=規格.金色, 粗體=True)
    管理工具.加入標題底線(masthead)

    badge = doc.add_paragraph()
    badge.paragraph_format.space_before = Pt(5)
    badge.paragraph_format.space_after = Pt(2)
    run = badge.add_run("SOP 01—12")
    規格.設定執行文字(run, 大小=11, 顏色=規格.次深藍, 粗體=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("智慧 5S\n標準作業程序")
    規格.設定執行文字(run, 大小=28, 顏色=規格.深藍, 粗體=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(15)
    run = subtitle.add_run("從正確登入到巡檢、改善、驗證、紅牌、離線補送與管理複盤")
    規格.設定執行文字(run, 大小=12.5, 顏色=規格.灰字)

    封面資料 = [
        ["管制項目", "內容", "管制項目", "內容"],
        ["文件編號", metadata.get("文件編號", "HXP-5S-SOP-001"), "文件版本", metadata.get("文件版本", "v1.0")],
        ["對應系統", metadata.get("對應系統", "智慧 5S 管理平台 v1.0.2"), "文件狀態", metadata.get("文件狀態", "正式版（核准後生效）")],
        ["制定單位", metadata.get("制定單位", "製造部"), "制定日期", metadata.get("制定日期", "2026 年 8 月 7 日")],
    ]
    規格.新增資料表(doc, 封面資料, 封面表格=True)

    if 核心要求:
        規格.新增提示框(doc, 核心要求, 封面=True)

    approval_label = doc.add_paragraph()
    approval_label.paragraph_format.space_before = Pt(6)
    approval_label.paragraph_format.space_after = Pt(5)
    run = approval_label.add_run("核准與生效")
    規格.設定執行文字(run, 大小=10.5, 顏色=規格.次深藍, 粗體=True)
    規格.新增資料表(doc, [
        ["編製", "審查", "核准", "生效日期"],
        ["製造部：____________", "單位主管：____________", "最高主管：____________", "______ 年 ____ 月 ____ 日"],
    ])

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("受控文件｜依畫面動作、完成訊號與異常停止點執行｜未核准不得擅自變更")
    規格.設定執行文字(run, 大小=9.5, 顏色=規格.灰字, 粗體=True)

    bottom = doc.add_paragraph()
    bottom.paragraph_format.space_before = Pt(18)
    bottom.paragraph_format.space_after = Pt(0)
    bottom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = bottom.add_run("化新精密有限公司｜製造部｜2026")
    規格.設定執行文字(run, 大小=9.5, 顏色=規格.灰字, 粗體=True)
    doc.add_page_break()


def 建立文件() -> Document:
    if not 來源檔案.exists():
        raise FileNotFoundError(f"找不到來源檔案：{來源檔案}")
    lines = 來源檔案.read_text(encoding="utf-8").splitlines()
    try:
        separator = lines.index("---")
    except ValueError as error:
        raise ValueError("來源文件缺少封面與本文分隔線 ---") from error

    doc = Document()
    doc.core_properties.title = "化新精密有限公司｜智慧 5S 標準作業程序（SOP）"
    doc.core_properties.subject = "智慧 5S 管理平台 v1.0.2 正式操作程序"
    doc.core_properties.author = "化新精密有限公司"
    doc.core_properties.last_modified_by = "化新精密有限公司"
    doc.core_properties.keywords = "智慧5S,SOP,巡檢,改善,驗證,紅牌,離線,LINE Bot"
    doc.core_properties.comments = "由受控 Markdown 來源自動產生；程序、畫面動作與系統版本同步。"

    規格.設定文件樣式(doc)
    section = doc.sections[0]
    規格.設定頁面(section)
    設定SOP頁首頁尾(section)
    新增SOP封面(doc, lines[:separator])
    規格.轉換本文(doc, lines[separator + 1:])
    return doc


def 擷取全文(path: Path) -> str:
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))
    return "".join(node.text or "" for node in document_xml.findall(".//w:t", namespaces))


def 稽核SOP(path: Path) -> dict:
    report = 規格.稽核結構(path)
    report["版型"] = "standard_business_brief"
    report["封面"] = "procedure_masthead"
    plain_text = 擷取全文(path)
    必備文字 = (
        "HXP-5S-SOP-001",
        "SOP-01",
        "SOP-12",
        "登入智慧 5S",
        "送出巡檢",
        "完成並送驗",
        "驗證通過",
        "駁回改善",
        "立即同步",
        "唯一 LINE Bot",
        "待改善",
        "改善中",
        "驗證中",
        "待處置",
        "處理中",
        "文件結束｜",
    )
    for text in 必備文字:
        if text not in plain_text:
            report["結構錯誤"].append(f"缺少 SOP 基準：{text}")
    report["結構錯誤數"] = len(report["結構錯誤"])
    report["結果"] = "通過" if not report["結構錯誤"] else "失敗"
    return report


def 主程式() -> int:
    doc = 建立文件()
    doc.save(輸出檔案)
    report = 稽核SOP(輸出檔案)
    稽核檔案.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print(f"已產生：{輸出檔案}")
    return 1 if report["結構錯誤數"] else 0


if __name__ == "__main__":
    sys.exit(主程式())
