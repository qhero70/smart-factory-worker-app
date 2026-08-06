#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
化新精密有限公司｜智慧 5S 完整操作手冊 Word 產生器

以受控 Markdown 建立正式 Word，沿用智慧 5S 文件設計系統、固定 DXA 表格、
真實 Word 清單、頁碼與結構稽核。文件文字只維護在智慧5S完整操作手冊_v1.0.md。
"""

from __future__ import annotations

import importlib.util
import json
import sys
import zipfile
from pathlib import Path
from typing import List, Sequence
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


目前目錄 = Path(__file__).resolve().parent
來源檔案 = 目前目錄 / "智慧5S完整操作手冊_v1.0.md"
輸出檔案 = 目前目錄 / "智慧5S完整操作手冊_v1.0.docx"
稽核檔案 = 目前目錄 / "智慧5S完整操作手冊_v1.0_結構稽核.json"
SOP產生器路徑 = 目前目錄 / "建立智慧5S_SOP.py"


def 載入SOP工具():
    if not SOP產生器路徑.exists():
        raise FileNotFoundError(f"找不到 SOP 產生器：{SOP產生器路徑}")
    spec = importlib.util.spec_from_file_location("智慧5S_SOP工具", SOP產生器路徑)
    if spec is None or spec.loader is None:
        raise ImportError(f"無法載入 SOP 產生器：{SOP產生器路徑}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


SOP工具 = 載入SOP工具()
規格 = SOP工具.規格
管理工具 = SOP工具.管理工具


def 選擇手冊欄寬(headers: Sequence[str]) -> List[int]:
    數量 = len(headers)
    joined = "|".join(headers)

    if 數量 == 2:
        if "不完整寫法" in joined:
            return [2800, 6560]
        if "畫面訊息" in joined or "問題" in joined:
            return [2900, 6460]
        return [2400, 6960]

    if 數量 == 3:
        if "正式責任" in joined:
            return [1500, 3500, 4360]
        if "錯誤風險" in joined:
            return [1700, 3600, 4060]
        if "系統結果" in joined:
            return [900, 1400, 3300, 3760]
        if "使用者處理" in joined:
            return [2200, 3200, 3960]
        if "建議使用者" in joined:
            return [1500, 4000, 3860]
        if "完成訊號" in joined:
            return [1500, 4000, 3860]
        return [2000, 3600, 3760]

    if 數量 == 4:
        if "管制項目" in joined:
            return [1500, 3180, 1500, 3180]
        if "編製" in joined and "核准" in joined:
            return [2340, 2340, 2340, 2340]
        if "版次" in joined:
            return [1000, 1500, 5060, 1800]
        if "我現在要做什麼" in joined:
            return [1800, 2200, 2200, 3160]
        if "不應執行" in joined:
            return [1500, 3300, 2100, 2460]
        if "系統行為" in joined:
            return [1400, 3100, 2100, 2760]
        if "管理輸出" in joined:
            return [1700, 4000, 3660]
        if "管理員查核" in joined:
            return [1800, 2300, 2300, 2960]
        if "回報項目" in joined:
            return [2200, 7160]
        return SOP工具.選擇SOP欄寬(headers)

    if 數量 == 5:
        if "必要證據" in joined:
            return [1200, 1700, 1900, 1700, 2860]
        return [1500, 1800, 1900, 1800, 2360]

    return SOP工具.選擇SOP欄寬(headers)


規格.選擇欄寬 = 選擇手冊欄寬


def 設定手冊頁首頁尾(section) -> None:
    頁首 = section.header.paragraphs[0]
    規格.清除段落(頁首)
    頁首.alignment = WD_ALIGN_PARAGRAPH.CENTER
    頁首.paragraph_format.space_after = Pt(0)
    run = 頁首.add_run("化新精密有限公司｜智慧 5S 管理平台 v1.0.2")
    規格.設定執行文字(run, 大小=8.5, 顏色=規格.灰字)

    首頁頁首 = section.first_page_header.paragraphs[0]
    規格.清除段落(首頁頁首)

    for footer in (section.footer, section.first_page_footer):
        頁尾 = footer.paragraphs[0]
        規格.清除段落(頁尾)
        頁尾.alignment = WD_ALIGN_PARAGRAPH.CENTER
        頁尾.paragraph_format.space_before = Pt(0)
        頁尾.paragraph_format.space_after = Pt(0)
        run = 頁尾.add_run("化新精密有限公司｜智慧 5S 完整操作手冊｜第 ")
        規格.設定執行文字(run, 大小=8.5, 顏色=規格.灰字)
        規格.加入頁碼欄位(頁尾)
        run = 頁尾.add_run(" 頁")
        規格.設定執行文字(run, 大小=8.5, 顏色=規格.灰字)


def 新增手冊封面(doc: Document, 前言: Sequence[str]) -> None:
    metadata, 核心要求 = 管理工具.解析封面資料(前言)

    masthead = doc.add_paragraph()
    masthead.paragraph_format.space_before = Pt(18)
    masthead.paragraph_format.space_after = Pt(14)
    run = masthead.add_run("化新精密有限公司｜正式使用者文件")
    規格.設定執行文字(run, 大小=11, 顏色=規格.金色, 粗體=True)
    管理工具.加入標題底線(masthead)

    badge = doc.add_paragraph()
    badge.paragraph_format.space_before = Pt(5)
    badge.paragraph_format.space_after = Pt(2)
    run = badge.add_run("USER MANUAL｜PWA 1.0.2")
    規格.設定執行文字(run, 大小=11, 顏色=規格.次深藍, 粗體=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("智慧 5S\n完整操作手冊")
    規格.設定執行文字(run, 大小=28, 顏色=規格.深藍, 粗體=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(15)
    run = subtitle.add_run("安裝、登入、巡檢、改善、驗證、盤點、紅牌、離線同步與主管複盤")
    規格.設定執行文字(run, 大小=12.5, 顏色=規格.灰字)

    封面資料 = [
        ["管制項目", "內容", "管制項目", "內容"],
        ["文件編號", metadata.get("文件編號", "HXP-5S-UM-001"), "文件版本", metadata.get("文件版本", "v1.0")],
        ["對應系統", metadata.get("對應系統", "智慧 5S 管理平台 v1.0.2"), "文件狀態", metadata.get("文件狀態", "正式版（核准後生效）")],
        ["制定單位", metadata.get("制定單位", "製造部"), "制定日期", metadata.get("制定日期", "2026 年 8 月 7 日")],
    ]
    規格.新增資料表(doc, 封面資料, 封面表格=True)

    if 核心要求:
        規格.新增提示框(doc, 核心要求, 封面=True)

    approval_label = doc.add_paragraph()
    approval_label.paragraph_format.space_before = Pt(6)
    approval_label.paragraph_format.space_after = Pt(5)
    run = approval_label.add_run("核准與生效")
    規格.設定執行文字(run, 大小=10.5, 顏色=規格.次深藍, 粗體=True)
    規格.新增資料表(doc, [
        ["編製", "審查", "核准", "生效日期"],
        ["製造部：____________", "單位主管：____________", "最高主管：____________", "______ 年 ____ 月 ____ 日"],
    ])

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("受控文件｜畫面名稱與按鈕須對應正式版本｜每次送出都確認單號、狀態與同步結果")
    規格.設定執行文字(run, 大小=9.5, 顏色=規格.灰字, 粗體=True)

    bottom = doc.add_paragraph()
    bottom.paragraph_format.space_before = Pt(18)
    bottom.paragraph_format.space_after = Pt(0)
    bottom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = bottom.add_run("化新精密有限公司｜製造部｜2026")
    規格.設定執行文字(run, 大小=9.5, 顏色=規格.灰字, 粗體=True)
    doc.add_page_break()


def 建立文件() -> Document:
    if not 來源檔案.exists():
        raise FileNotFoundError(f"找不到來源檔案：{來源檔案}")
    lines = 來源檔案.read_text(encoding="utf-8").splitlines()
    try:
        separator = lines.index("---")
    except ValueError as error:
        raise ValueError("來源文件缺少封面與本文分隔線 ---") from error

    doc = Document()
    doc.core_properties.title = "化新精密有限公司｜智慧 5S 完整操作手冊"
    doc.core_properties.subject = "智慧 5S 管理平台 v1.0.2 正式使用者操作手冊"
    doc.core_properties.author = "化新精密有限公司"
    doc.core_properties.last_modified_by = "化新精密有限公司"
    doc.core_properties.keywords = "智慧5S,操作手冊,PWA,巡檢,改善,驗證,紅牌,離線同步"
    doc.core_properties.comments = "由受控 Markdown 來源自動產生；畫面、流程、文件與系統版本同步。"

    規格.設定文件樣式(doc)
    section = doc.sections[0]
    規格.設定頁面(section)
    設定手冊頁首頁尾(section)
    新增手冊封面(doc, lines[:separator])
    規格.轉換本文(doc, lines[separator + 1:])
    return doc


def 擷取全文(path: Path) -> str:
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))
    return "".join(node.text or "" for node in document_xml.findall(".//w:t", namespaces))


def 稽核操作手冊(path: Path) -> dict:
    report = 規格.稽核結構(path)
    report["版型"] = "standard_business_brief"
    report["封面"] = "user_manual_masthead"
    plain_text = 擷取全文(path)
    必備文字 = (
        "HXP-5S-UM-001",
        "一、文件管制與使用方式",
        "二十二、常見問答",
        "登入智慧 5S",
        "立即開始巡檢",
        "送出巡檢",
        "完成並送驗",
        "驗證通過",
        "駁回改善",
        "新增物品盤點",
        "完成並結案",
        "立即同步",
        "清除快取並更新",
        "待同步 0",
        "唯一 LINE Bot",
        "文件結束｜",
    )
    for text in 必備文字:
        if text not in plain_text:
            report["結構錯誤"].append(f"缺少操作手冊基準：{text}")
    if report.get("表格數", 0) < 35:
        report["結構錯誤"].append("操作手冊表格數不足 35，疑似內容未完整輸出")
    if report.get("真實編號定義數", 0) < 35:
        report["結構錯誤"].append("操作手冊真實編號定義不足 35，疑似步驟未完整輸出")
    report["結構錯誤數"] = len(report["結構錯誤"])
    report["結果"] = "通過" if not report["結構錯誤"] else "失敗"
    return report


def 主程式() -> int:
    doc = 建立文件()
    doc.save(輸出檔案)
    report = 稽核操作手冊(輸出檔案)
    稽核檔案.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print(f"已產生：{輸出檔案}")
    return 1 if report["結構錯誤數"] else 0


if __name__ == "__main__":
    sys.exit(主程式())
