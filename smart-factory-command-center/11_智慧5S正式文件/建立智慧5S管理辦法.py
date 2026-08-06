#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
化新精密有限公司｜智慧 5S 管理辦法 Word 產生器

用途：
1. 以「智慧5S管理辦法_v1.0.md」作為唯一文字來源。
2. 沿用智慧 5S 正式文件的 standard_business_brief 版型。
3. 使用 memo_masthead 受控文件封面、真實 Word 清單與固定 DXA 表格。
4. 產生後執行結構稽核，確保文件、表格與重要制度基準完整。

執行方式：
  $CODEX_PRIMARY_RUNTIME_PYTHON 建立智慧5S管理辦法.py
"""

from __future__ import annotations

import importlib.util
import json
import sys
import zipfile
from pathlib import Path
from typing import List, Sequence
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


目前目錄 = Path(__file__).resolve().parent
來源檔案 = 目前目錄 / "智慧5S管理辦法_v1.0.md"
輸出檔案 = 目前目錄 / "智慧5S管理辦法_v1.0.docx"
稽核檔案 = 目前目錄 / "智慧5S管理辦法_v1.0_結構稽核.json"
共用產生器路徑 = 目前目錄 / "建立智慧5S推動計畫書.py"


def 載入共用規格():
    """從既有推動計畫書產生器載入同一份文件設計系統。"""
    if not 共用產生器路徑.exists():
        raise FileNotFoundError(f"找不到共用產生器：{共用產生器路徑}")
    spec = importlib.util.spec_from_file_location("智慧5S文件共用規格", 共用產生器路徑)
    if spec is None or spec.loader is None:
        raise ImportError(f"無法載入共用產生器：{共用產生器路徑}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


規格 = 載入共用規格()
原選擇欄寬 = 規格.選擇欄寬


def 選擇管理辦法欄寬(headers: Sequence[str]) -> List[int]:
    """依管理辦法的表格語意調整欄寬，合計固定為 9,360 DXA。"""
    數量 = len(headers)
    joined = "|".join(headers)

    if 數量 == 2:
        if "正式基準" in joined:
            return [2400, 6960]
        return [2200, 7160]

    if 數量 == 3:
        if "角色" in joined and "權限邊界" in joined:
            return [1650, 4550, 3160]
        if "類別" in joined and "現場判定重點" in joined:
            return [1300, 3400, 4660]
        if "狀態" in joined and "責任動作" in joined:
            return [1500, 2860, 5000]
        if "資料類型" in joined and "受控分頁" in joined:
            return [1700, 4360, 3300]
        return [1800, 3780, 3780]

    if 數量 == 4:
        if "管制項目" in joined:
            return [1500, 3180, 1500, 3180]
        if "編製" in joined and "核准" in joined:
            return [2340, 2340, 2340, 2340]
        if "版次" in joined:
            return [1000, 1500, 5060, 1800]
        if "執行對象" in joined and "最低頻率" in joined:
            return [1800, 1500, 3260, 2800]
        if "指標" in joined and "管理意義" in joined:
            return [1750, 3050, 1400, 3160]
        return 原選擇欄寬(headers)

    if 數量 == 5 and "分數" in joined:
        return [800, 1400, 1400, 3000, 2760]

    return 原選擇欄寬(headers)


# 共用 Markdown 轉換器內的表格函式會從模組全域取用欄寬函式。
規格.選擇欄寬 = 選擇管理辦法欄寬


def 解析封面資料(前言: Sequence[str]) -> tuple[dict[str, str], str]:
    metadata: dict[str, str] = {}
    核心要求 = ""
    for line in 前言:
        stripped = line.strip().rstrip("  ")
        if stripped.startswith(">"):
            核心要求 = stripped[1:].strip()
        elif "：" in stripped and not stripped.startswith("#"):
            key, value = stripped.split("：", 1)
            metadata[key.strip()] = value.strip()
    return metadata, 核心要求


def 設定管理辦法頁首頁尾(section) -> None:
    頁首 = section.header.paragraphs[0]
    規格.清除段落(頁首)
    頁首.alignment = WD_ALIGN_PARAGRAPH.CENTER
    頁首.paragraph_format.space_after = Pt(0)
    run = 頁首.add_run("化新精密有限公司｜智慧 5S 管理平台 v1.0.2")
    規格.設定執行文字(run, 大小=8.5, 顏色=規格.灰字)

    首頁頁首 = section.first_page_header.paragraphs[0]
    規格.清除段落(首頁頁首)

    for footer in (section.footer, section.first_page_footer):
        頁尾 = footer.paragraphs[0]
        規格.清除段落(頁尾)
        頁尾.alignment = WD_ALIGN_PARAGRAPH.CENTER
        頁尾.paragraph_format.space_before = Pt(0)
        頁尾.paragraph_format.space_after = Pt(0)
        run = 頁尾.add_run("化新精密有限公司｜智慧 5S 管理辦法｜第 ")
        規格.設定執行文字(run, 大小=8.5, 顏色=規格.灰字)
        規格.加入頁碼欄位(頁尾)
        run = 頁尾.add_run(" 頁")
        規格.設定執行文字(run, 大小=8.5, 顏色=規格.灰字)


def 加入標題底線(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), 規格.金色)
    pbdr.append(bottom)


def 新增管理辦法封面(doc: Document, 前言: Sequence[str]) -> None:
    metadata, 核心要求 = 解析封面資料(前言)

    masthead = doc.add_paragraph()
    masthead.paragraph_format.space_before = Pt(18)
    masthead.paragraph_format.space_after = Pt(14)
    masthead.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = masthead.add_run("化新精密有限公司｜內部管理制度")
    規格.設定執行文字(run, 大小=11, 顏色=規格.金色, 粗體=True)
    加入標題底線(masthead)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("智慧 5S\n管理辦法")
    規格.設定執行文字(run, 大小=29, 顏色=規格.深藍, 粗體=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run("巡檢、改善、紅牌、驗證、通知與資料治理正式規範")
    規格.設定執行文字(run, 大小=13, 顏色=規格.灰字)

    封面資料 = [
        ["管制項目", "內容", "管制項目", "內容"],
        ["文件編號", metadata.get("文件編號", "HXP-5S-MGT-001"), "文件版本", metadata.get("文件版本", "v1.0")],
        ["對應系統", metadata.get("對應系統", "智慧 5S 管理平台 v1.0.2"), "文件狀態", metadata.get("文件狀態", "正式版（核准後生效）")],
        ["制定單位", metadata.get("制定單位", "製造部"), "制定日期", metadata.get("制定日期", "2026 年 8 月 7 日")],
    ]
    規格.新增資料表(doc, 封面資料, 封面表格=True)

    if 核心要求:
        規格.新增提示框(doc, 核心要求, 封面=True)

    approval_label = doc.add_paragraph()
    approval_label.paragraph_format.space_before = Pt(6)
    approval_label.paragraph_format.space_after = Pt(5)
    run = approval_label.add_run("核准與生效")
    規格.設定執行文字(run, 大小=10.5, 顏色=規格.次深藍, 粗體=True)

    核准資料 = [
        ["編製", "審查", "核准", "生效日期"],
        ["製造部：____________", "單位主管：____________", "最高主管：____________", "______ 年 ____ 月 ____ 日"],
    ]
    規格.新增資料表(doc, 核准資料)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run("受控文件｜未經核准不得擅自變更｜歷史版本不覆蓋")
    規格.設定執行文字(run, 大小=9.5, 顏色=規格.灰字, 粗體=True)

    bottom = doc.add_paragraph()
    bottom.paragraph_format.space_before = Pt(20)
    bottom.paragraph_format.space_after = Pt(0)
    bottom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = bottom.add_run("化新精密有限公司｜製造部｜2026")
    規格.設定執行文字(run, 大小=9.5, 顏色=規格.灰字, 粗體=True)

    doc.add_page_break()


def 建立文件() -> Document:
    if not 來源檔案.exists():
        raise FileNotFoundError(f"找不到來源檔案：{來源檔案}")

    markdown = 來源檔案.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    try:
        separator = lines.index("---")
    except ValueError as error:
        raise ValueError("來源文件缺少封面與本文分隔線 ---") from error

    doc = Document()
    doc.core_properties.title = "化新精密有限公司｜智慧 5S 管理辦法"
    doc.core_properties.subject = "智慧 5S 管理平台 v1.0.2 正式管理制度"
    doc.core_properties.author = "化新精密有限公司"
    doc.core_properties.last_modified_by = "化新精密有限公司"
    doc.core_properties.keywords = "智慧5S,管理辦法,巡檢,改善,紅牌,PWA,GAS,LINE Bot,資料治理"
    doc.core_properties.comments = "由受控 Markdown 來源自動產生；程式、文件與系統版本同步。"

    規格.設定文件樣式(doc)
    section = doc.sections[0]
    規格.設定頁面(section)
    設定管理辦法頁首頁尾(section)
    新增管理辦法封面(doc, lines[:separator])
    規格.轉換本文(doc, lines[separator + 1:])
    return doc


def 擷取全文(path: Path) -> str:
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))
    return "".join(node.text or "" for node in document_xml.findall(".//w:t", namespaces))


def 稽核管理辦法(path: Path) -> dict:
    report = 規格.稽核結構(path)
    report["版型"] = "standard_business_brief"
    report["封面"] = "memo_masthead"

    plain_text = 擷取全文(path)
    必備文字 = (
        "HXP-5S-MGT-001",
        "改善單正式預設期限為建立日起 7 日",
        "紅牌預設處置期限為掛牌日起 30 日",
        "待改善",
        "改善中",
        "驗證中",
        "已結案",
        "待處置",
        "唯一 LINE Bot",
        "文件結束｜",
    )
    for text in 必備文字:
        if text not in plain_text:
            report["結構錯誤"].append(f"缺少管理基準：{text}")

    report["結構錯誤數"] = len(report["結構錯誤"])
    report["結果"] = "通過" if not report["結構錯誤"] else "失敗"
    return report


def 主程式() -> int:
    doc = 建立文件()
    doc.save(輸出檔案)
    report = 稽核管理辦法(輸出檔案)
    稽核檔案.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print(f"已產生：{輸出檔案}")
    if report["結構錯誤數"]:
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(主程式())
