#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
化新精密有限公司｜智慧 5S 推動計畫書 Word 產生器

用途：
1. 以「智慧5S推動計畫書_v1.0.md」作為唯一文字來源。
2. 產生可直接送審的繁體中文 Word 文件。
3. 套用 standard_business_brief 版型與 proposal_centerpiece 封面。
4. 使用真正的 Word 標題、編號、頁碼與固定 DXA 表格幾何。
5. 產生後執行結構稽核，避免表格、版面或假清單格式漂移。

執行方式：
  $CODEX_PRIMARY_RUNTIME_PYTHON 建立智慧5S推動計畫書.py
"""

from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path
from typing import Iterable, List, Sequence
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


目前目錄 = Path(__file__).resolve().parent
來源檔案 = 目前目錄 / "智慧5S推動計畫書_v1.0.md"
輸出檔案 = 目前目錄 / "智慧5S推動計畫書_v1.0.docx"
稽核檔案 = 目前目錄 / "智慧5S推動計畫書_v1.0_結構稽核.json"


# standard_business_brief 精確版型參數
頁面寬度_DXA = 12240
頁面高度_DXA = 15840
邊界_DXA = 1440
內容寬度_DXA = 9360
表格縮排_DXA = 120
儲存格上_DXA = 80
儲存格下_DXA = 80
儲存格左_DXA = 120
儲存格右_DXA = 120

西文字型 = "Calibri"
# 使用可合法散布的 Noto Sans CJK TC；未安裝時 Word 會自動改用系統繁中字型。
中文字型 = "Noto Sans CJK TC"
標題藍 = "2E74B5"
深藍 = "0B2545"
次深藍 = "1F4D78"
灰字 = "5B6573"
淡灰 = "F2F4F7"
淡藍灰 = "E8EEF5"
提示底色 = "F4F6F9"
金色 = "B38B2E"
白色 = "FFFFFF"
框線灰 = "B8C2CC"


def 設定執行文字(run, *, 大小: float | None = None, 顏色: str | None = None,
             粗體: bool | None = None, 斜體: bool | None = None) -> None:
    """同時設定西文與東亞字型，確保繁體中文在 Word 中正確顯示。"""
    run.font.name = 西文字型
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), 西文字型)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), 西文字型)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), 中文字型)
    if 大小 is not None:
        run.font.size = Pt(大小)
    if 顏色 is not None:
        run.font.color.rgb = RGBColor.from_string(顏色)
    if 粗體 is not None:
        run.bold = 粗體
    if 斜體 is not None:
        run.italic = 斜體


def 設定樣式字型(style, *, 大小: float, 顏色: str = "000000", 粗體: bool = False) -> None:
    style.font.name = 西文字型
    style.font.size = Pt(大小)
    style.font.color.rgb = RGBColor.from_string(顏色)
    style.font.bold = 粗體
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), 西文字型)
    rfonts.set(qn("w:hAnsi"), 西文字型)
    rfonts.set(qn("w:eastAsia"), 中文字型)


def 設定文件樣式(doc: Document) -> None:
    styles = doc.styles

    正文 = styles["Normal"]
    設定樣式字型(正文, 大小=11)
    正文.paragraph_format.space_before = Pt(0)
    正文.paragraph_format.space_after = Pt(6)
    正文.paragraph_format.line_spacing = 1.10

    標題一 = styles["Heading 1"]
    設定樣式字型(標題一, 大小=16, 顏色=標題藍, 粗體=True)
    標題一.paragraph_format.space_before = Pt(16)
    標題一.paragraph_format.space_after = Pt(8)
    標題一.paragraph_format.line_spacing = 1.0
    標題一.paragraph_format.keep_with_next = True
    標題一.paragraph_format.keep_together = True

    標題二 = styles["Heading 2"]
    設定樣式字型(標題二, 大小=13, 顏色=標題藍, 粗體=True)
    標題二.paragraph_format.space_before = Pt(12)
    標題二.paragraph_format.space_after = Pt(6)
    標題二.paragraph_format.line_spacing = 1.0
    標題二.paragraph_format.keep_with_next = True
    標題二.paragraph_format.keep_together = True

    標題三 = styles["Heading 3"]
    設定樣式字型(標題三, 大小=12, 顏色=次深藍, 粗體=True)
    標題三.paragraph_format.space_before = Pt(8)
    標題三.paragraph_format.space_after = Pt(4)
    標題三.paragraph_format.line_spacing = 1.0
    標題三.paragraph_format.keep_with_next = True
    標題三.paragraph_format.keep_together = True


def 設定頁面(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def 清除段落(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def 加入頁碼欄位(paragraph) -> None:
    run = paragraph.add_run()
    設定執行文字(run, 大小=8.5, 顏色=灰字)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    visible = OxmlElement("w:t")
    visible.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, visible, end])


def 設定頁首頁尾(section) -> None:
    頁首 = section.header.paragraphs[0]
    清除段落(頁首)
    頁首.alignment = WD_ALIGN_PARAGRAPH.CENTER
    頁首.paragraph_format.space_after = Pt(0)
    run = 頁首.add_run("化新精密有限公司｜智慧 5S 管理平台 v1.0.2")
    設定執行文字(run, 大小=8.5, 顏色=灰字)

    首頁頁首 = section.first_page_header.paragraphs[0]
    清除段落(首頁頁首)

    for footer in (section.footer, section.first_page_footer):
        頁尾 = footer.paragraphs[0]
        清除段落(頁尾)
        頁尾.alignment = WD_ALIGN_PARAGRAPH.CENTER
        頁尾.paragraph_format.space_before = Pt(0)
        頁尾.paragraph_format.space_after = Pt(0)
        run = 頁尾.add_run("化新精密有限公司｜智慧 5S 推動計畫書｜第 ")
        設定執行文字(run, 大小=8.5, 顏色=灰字)
        加入頁碼欄位(頁尾)
        run = 頁尾.add_run(" 頁")
        設定執行文字(run, 大小=8.5, 顏色=灰字)


def 設定儲存格底色(cell, 色碼: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), 色碼)
    shd.set(qn("w:val"), "clear")


def 設定表格框線(table, 色碼: str = 框線灰, 粗細: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(粗細))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), 色碼)


def 設定儲存格邊界(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (
        ("top", 儲存格上_DXA),
        ("bottom", 儲存格下_DXA),
        ("start", 儲存格左_DXA),
        ("end", 儲存格右_DXA),
    ):
        tag = qn(f"w:{name}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def 設定列不拆頁(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def 設定表頭重複(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def 設定表格幾何(table, 欄寬_DXA: Sequence[int]) -> None:
    if sum(欄寬_DXA) != 內容寬度_DXA:
        raise ValueError(f"表格欄寬合計必須為 {內容寬度_DXA} DXA：{欄寬_DXA}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(內容寬度_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(表格縮排_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in 欄寬_DXA:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        設定列不拆頁(row)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(欄寬_DXA[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(欄寬_DXA[index]))
            tc_w.set(qn("w:type"), "dxa")
            設定儲存格邊界(cell)


def 選擇欄寬(headers: Sequence[str]) -> List[int]:
    數量 = len(headers)
    joined = "|".join(headers)
    if 數量 == 1:
        return [9360]
    if 數量 == 2:
        return [2700, 6660]
    if 數量 == 3:
        return [1800, 3780, 3780]
    if 數量 == 4:
        if "管制項目" in joined:
            return [1500, 3180, 1500, 3180]
        if "編製" in joined and "核准" in joined:
            return [2340, 2340, 2340, 2340]
        if "對象" in joined and "建議時間" in joined:
            return [1500, 3500, 1600, 2760]
        return [1500, 2860, 2860, 2140]
    if 數量 == 5:
        if "階段" in joined and "階段出口" in joined:
            return [1500, 1500, 1700, 2860, 1800]
        if "風險" in joined:
            return [1100, 1700, 2500, 2460, 1600]
        if "指標" in joined and "檢視頻率" in joined:
            return [1600, 2200, 1600, 2460, 1500]
        return [1400, 1900, 2400, 2160, 1500]
    raise ValueError(f"尚未定義 {數量} 欄表格的欄寬：{headers}")


def 設定段落行距(paragraph, *, after: float = 0, line: float = 1.0) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def 加入超連結(paragraph, 顯示文字: str, 網址: str) -> None:
    relation_id = paragraph.part.relate_to(網址, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), 標題藍)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), 西文字型)
    fonts.set(qn("w:hAnsi"), 西文字型)
    fonts.set(qn("w:eastAsia"), 中文字型)
    rpr.extend([fonts, color, underline])
    text = OxmlElement("w:t")
    text.text = 顯示文字
    run.extend([rpr, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def 加入格式化文字(paragraph, 文字: str, *, 大小: float = 11,
              顏色: str = "000000", 粗體: bool = False) -> None:
    token = re.compile(r"(`[^`]+`|https?://[^\s<>]+)")
    位置 = 0
    for match in token.finditer(文字):
        if match.start() > 位置:
            run = paragraph.add_run(文字[位置:match.start()])
            設定執行文字(run, 大小=大小, 顏色=顏色, 粗體=粗體)
        片段 = match.group(0)
        if 片段.startswith("http"):
            加入超連結(paragraph, 片段, 片段)
        else:
            run = paragraph.add_run(片段[1:-1])
            設定執行文字(run, 大小=max(大小 - 0.5, 8), 顏色=次深藍, 粗體=True)
        位置 = match.end()
    if 位置 < len(文字):
        run = paragraph.add_run(文字[位置:])
        設定執行文字(run, 大小=大小, 顏色=顏色, 粗體=粗體)


def 清理表格文字(文字: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", 文字.strip())


def 新增資料表(doc: Document, 資料: Sequence[Sequence[str]], *, 封面表格: bool = False) -> None:
    if not 資料:
        return
    欄寬 = 選擇欄寬(資料[0])
    table = doc.add_table(rows=len(資料), cols=len(資料[0]))
    設定表格幾何(table, 欄寬)
    設定表格框線(table, 框線灰, 6)
    設定表頭重複(table.rows[0])

    for row_index, row in enumerate(資料):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            paragraph = cell.paragraphs[0]
            清除段落(paragraph)
            設定段落行距(paragraph, after=0, line=1.0)

            is_header = row_index == 0
            is_label = 封面表格 and col_index % 2 == 0
            if is_header:
                設定儲存格底色(cell, 淡灰)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif is_label:
                設定儲存格底色(cell, 淡藍灰)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
                )

            run = paragraph.add_run(清理表格文字(value))
            設定執行文字(
                run,
                大小=9.3 if len(資料[0]) >= 5 else 9.8,
                顏色=深藍 if (is_header or is_label) else "000000",
                粗體=is_header or is_label,
            )

    spacing = doc.add_paragraph()
    spacing.paragraph_format.space_after = Pt(2)


def 新增提示框(doc: Document, 文字: str, *, 封面: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    設定表格幾何(table, [內容寬度_DXA])
    設定表格框線(table, 深藍 if 封面 else 框線灰, 8)
    cell = table.cell(0, 0)
    設定儲存格底色(cell, 深藍 if 封面 else 提示底色)
    paragraph = cell.paragraphs[0]
    清除段落(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if 封面 else WD_ALIGN_PARAGRAPH.LEFT
    設定段落行距(paragraph, after=0, line=1.15)
    run = paragraph.add_run(文字)
    設定執行文字(
        run,
        大小=11.5 if 封面 else 10.5,
        顏色=白色 if 封面 else 深藍,
        粗體=True,
    )
    spacing = doc.add_paragraph()
    spacing.paragraph_format.space_after = Pt(1)


def 下一個識別碼(root, tag_name: str, attr_name: str) -> int:
    values = []
    for node in root.findall(qn(tag_name)):
        raw = node.get(qn(attr_name))
        if raw is not None and raw.isdigit():
            values.append(int(raw))
    return max(values, default=0) + 1


def 建立清單編號(doc: Document, 類型: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = 下一個識別碼(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = 下一個識別碼(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if 類型 == "項目" else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "●" if 類型 == "項目" else "%1.")
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, indent, spacing])
    level.extend([start, num_fmt, level_text, level_jc, ppr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def 套用清單編號(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def 新增清單項目(doc: Document, 文字: str, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    套用清單編號(paragraph, num_id)
    加入格式化文字(paragraph, 文字)


def 新增核取項目(doc: Document, 文字: str, 已完成: bool) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    label = "【完成】" if 已完成 else "【待辦】"
    run = paragraph.add_run(label + " ")
    設定執行文字(run, 大小=10.5, 顏色=次深藍 if 已完成 else "7A5A00", 粗體=True)
    加入格式化文字(paragraph, 文字)


def 新增一般段落(doc: Document, 文字: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.widow_control = True
    加入格式化文字(paragraph, 文字)


def 新增封面(doc: Document, 前言: Sequence[str]) -> None:
    metadata = {}
    核心目標 = ""
    for line in 前言:
        stripped = line.strip().rstrip("  ")
        if stripped.startswith(">"):
            核心目標 = stripped[1:].strip()
        elif "：" in stripped and not stripped.startswith("#"):
            key, value = stripped.split("：", 1)
            metadata[key.strip()] = value.strip()

    top = doc.add_paragraph()
    top.paragraph_format.space_before = Pt(36)
    top.paragraph_format.space_after = Pt(8)
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top.add_run("化新精密有限公司")
    設定執行文字(run, 大小=13, 顏色=灰字, 粗體=True)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(18)
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("智慧製造・現場管理・持續改善")
    設定執行文字(run, 大小=10.5, 顏色=金色, 粗體=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("智慧 5S\n推動計畫書")
    設定執行文字(run, 大小=30, 顏色=深藍, 粗體=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("從 A9 示範區建立可複製、可追溯、可持續的現場管理閉環")
    設定執行文字(run, 大小=13.5, 顏色=灰字)

    cover_data = [
        ["文件版本", "內容", "管制項目", "內容"],
        ["文件版本", metadata.get("文件版本", "v1.0"), "系統版本", metadata.get("對應系統", "智慧 5S 管理平台 v1.0.2")],
        ["文件狀態", metadata.get("文件狀態", "正式推動版"), "編製日期", metadata.get("編製日期", "2026 年 8 月 6 日")],
        ["主責單位", metadata.get("主責單位", "製造部"), "示範區", "A9 四個示範區"],
    ]
    新增資料表(doc, cover_data, 封面表格=True)

    if 核心目標:
        新增提示框(doc, 核心目標, 封面=True)

    bottom = doc.add_paragraph()
    bottom.paragraph_format.space_before = Pt(28)
    bottom.paragraph_format.space_after = Pt(0)
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = bottom.add_run("化新精密有限公司｜製造部｜2026")
    設定執行文字(run, 大小=10, 顏色=灰字, 粗體=True)

    doc.add_page_break()


def 是表格分隔列(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def 解析表格(lines: Sequence[str], start: int) -> tuple[List[List[str]], int]:
    rows: List[List[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        if not 是表格分隔列(line):
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
        index += 1
    if rows:
        width = len(rows[0])
        if any(len(row) != width for row in rows):
            raise ValueError(f"Markdown 表格欄數不一致，起始列：{start + 1}")
    return rows, index


def 轉換本文(doc: Document, lines: Sequence[str]) -> None:
    index = 0
    清單類型 = None
    清單識別碼 = None

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if not line:
            清單類型 = None
            清單識別碼 = None
            index += 1
            continue

        if line == "<!-- 換頁 -->":
            doc.add_page_break()
            清單類型 = None
            清單識別碼 = None
            index += 1
            continue

        if line == "---":
            index += 1
            continue

        if line.startswith("|"):
            table_rows, index = 解析表格(lines, index)
            新增資料表(doc, table_rows)
            清單類型 = None
            清單識別碼 = None
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            doc.add_paragraph(heading.group(2).strip(), style=f"Heading {level}")
            清單類型 = None
            清單識別碼 = None
            index += 1
            continue

        if line.startswith(">"):
            新增提示框(doc, line[1:].strip())
            清單類型 = None
            清單識別碼 = None
            index += 1
            continue

        checkbox = re.match(r"^- \[([xX ])\]\s+(.+)$", line)
        if checkbox:
            新增核取項目(doc, checkbox.group(2), checkbox.group(1).lower() == "x")
            清單類型 = None
            清單識別碼 = None
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            if 清單類型 != "項目" or 清單識別碼 is None:
                清單類型 = "項目"
                清單識別碼 = 建立清單編號(doc, "項目")
            新增清單項目(doc, bullet.group(1), 清單識別碼)
            index += 1
            continue

        decimal = re.match(r"^\d+\.\s+(.+)$", line)
        if decimal:
            if 清單類型 != "編號" or 清單識別碼 is None:
                清單類型 = "編號"
                清單識別碼 = 建立清單編號(doc, "編號")
            新增清單項目(doc, decimal.group(1), 清單識別碼)
            index += 1
            continue

        if line.startswith("文件結束｜"):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(line)
            設定執行文字(run, 大小=9, 顏色=灰字)
        else:
            新增一般段落(doc, line)

        清單類型 = None
        清單識別碼 = None
        index += 1


def 建立文件() -> Document:
    if not 來源檔案.exists():
        raise FileNotFoundError(f"找不到來源檔案：{來源檔案}")

    markdown = 來源檔案.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    try:
        separator = lines.index("---")
    except ValueError as error:
        raise ValueError("來源文件缺少封面與本文分隔線 ---") from error

    doc = Document()
    doc.core_properties.title = "化新精密有限公司｜智慧 5S 推動計畫書"
    doc.core_properties.subject = "智慧 5S 管理平台 v1.0.2 正式推動計畫"
    doc.core_properties.author = "化新精密有限公司"
    doc.core_properties.last_modified_by = "化新精密有限公司"
    doc.core_properties.keywords = "智慧5S,PWA,GAS,LINE Bot,現場管理,持續改善"
    doc.core_properties.comments = "由受控 Markdown 來源自動產生；程式、文件與系統版本同步。"

    設定文件樣式(doc)
    section = doc.sections[0]
    設定頁面(section)
    設定頁首頁尾(section)
    新增封面(doc, lines[:separator])
    轉換本文(doc, lines[separator + 1:])
    return doc


def 稽核結構(path: Path) -> dict:
    namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        document_xml = ET.fromstring(archive.read("word/document.xml"))
        styles_xml = ET.fromstring(archive.read("word/styles.xml"))
        numbering_xml = ET.fromstring(archive.read("word/numbering.xml"))

    errors: List[str] = []
    tables = document_xml.findall(".//w:tbl", namespaces)
    for table_index, table in enumerate(tables, start=1):
        tbl_w = table.find("./w:tblPr/w:tblW", namespaces)
        tbl_ind = table.find("./w:tblPr/w:tblInd", namespaces)
        layout = table.find("./w:tblPr/w:tblLayout", namespaces)
        if tbl_w is None or tbl_w.get(qn("w:type")) != "dxa" or tbl_w.get(qn("w:w")) != str(內容寬度_DXA):
            errors.append(f"第 {table_index} 個表格 tblW 不符合 {內容寬度_DXA} DXA")
        if tbl_ind is None or tbl_ind.get(qn("w:type")) != "dxa" or tbl_ind.get(qn("w:w")) != str(表格縮排_DXA):
            errors.append(f"第 {table_index} 個表格 tblInd 不符合 {表格縮排_DXA} DXA")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            errors.append(f"第 {table_index} 個表格未使用 fixed layout")

        grid_widths = [int(node.get(qn("w:w"), "0")) for node in table.findall("./w:tblGrid/w:gridCol", namespaces)]
        if sum(grid_widths) != 內容寬度_DXA:
            errors.append(f"第 {table_index} 個表格 tblGrid 合計不是 {內容寬度_DXA} DXA")

        for row_index, row in enumerate(table.findall("./w:tr", namespaces), start=1):
            widths = []
            for cell in row.findall("./w:tc", namespaces):
                tc_w = cell.find("./w:tcPr/w:tcW", namespaces)
                if tc_w is None or tc_w.get(qn("w:type")) != "dxa":
                    widths.append(0)
                else:
                    widths.append(int(tc_w.get(qn("w:w"), "0")))
            if sum(widths) != 內容寬度_DXA:
                errors.append(f"第 {table_index} 個表格第 {row_index} 列 tcW 合計不是 {內容寬度_DXA} DXA")

    section = document_xml.find(".//w:sectPr", namespaces)
    page_size = section.find("./w:pgSz", namespaces) if section is not None else None
    page_margin = section.find("./w:pgMar", namespaces) if section is not None else None
    if page_size is None or page_size.get(qn("w:w")) != str(頁面寬度_DXA) or page_size.get(qn("w:h")) != str(頁面高度_DXA):
        errors.append("頁面尺寸不是 US Letter 直式")
    if page_margin is None or any(page_margin.get(qn(f"w:{side}")) != str(邊界_DXA) for side in ("top", "right", "bottom", "left")):
        errors.append("頁面四邊界不是 1 英吋")

    style_ids = {node.get(qn("w:styleId")) for node in styles_xml.findall("./w:style", namespaces)}
    for style_id in ("Normal", "Heading1", "Heading2", "Heading3"):
        if style_id not in style_ids:
            errors.append(f"缺少 Word 樣式：{style_id}")

    num_count = len(numbering_xml.findall("./w:num", namespaces))
    fake_list_count = 0
    for paragraph in document_xml.findall(".//w:body/w:p", namespaces):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", namespaces)).strip()
        if re.match(r"^(- |\d+\. )", text):
            fake_list_count += 1
    if fake_list_count:
        errors.append(f"發現 {fake_list_count} 個假清單段落")

    plain_text = "".join(node.text or "" for node in document_xml.findall(".//w:t", namespaces))
    for forbidden in ("<!--", "turn0", "TODO", "請在此處自行"):
        if forbidden in plain_text:
            errors.append(f"文件含有不應交付的文字：{forbidden}")

    return {
        "文件": path.name,
        "版型": "standard_business_brief",
        "封面": "proposal_centerpiece",
        "表格數": len(tables),
        "真實編號定義數": num_count,
        "假清單數": fake_list_count,
        "結構錯誤數": len(errors),
        "結構錯誤": errors,
        "結果": "通過" if not errors else "失敗",
    }


def 主程式() -> int:
    doc = 建立文件()
    doc.save(輸出檔案)
    report = 稽核結構(輸出檔案)
    稽核檔案.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print(f"已產生：{輸出檔案}")
    if report["結構錯誤數"]:
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(主程式())
